import docx

doc = docx.Document('probe.docx')
doc.paragraphs[0].runs[0].text = doc.paragraphs[0].runs[0].text.replace('ФИО', 'ТЕСТ')
doc.save('probe_done.docx')

